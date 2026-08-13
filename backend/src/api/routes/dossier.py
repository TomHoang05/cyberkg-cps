"""Dossier endpoint — Sprint 4 / T062.

GET /api/v1/attacks/{slug}/dossier

Two rendering paths:
  1. LLM narrative (fast) — if POST /attacks/{slug}/narrative was called first,
     the stored prose is formatted into a .docx and returned instantly.
  2. KG-structured (fallback) — if no narrative exists, the endpoint queries
     Q1–Q4 directly and builds a structured teaching document with no LLM call.
     This path requires no API keys and always works when Neo4j is running.

Query params
------------
  audience        : "instructor" (default) | "student"
  deployment_size : "standard" (default) | "large"

Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document
Content-Disposition: attachment; filename="{slug}_dossier.docx"
"""
from __future__ import annotations

import io
from datetime import datetime, timezone
from typing import Literal

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from neo4j import Driver

from src.api.attack_id_map import resolve_attack_id
from src.api.dependencies import get_driver
from src.api.services.narrative_service import get_narrative
from src.kg.queries import q1_surface, q2_chain, q3_consequence, q4_roles

router = APIRouter(prefix="/attacks", tags=["dossier"])

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# ---------------------------------------------------------------------------
# Endpoint
# ---------------------------------------------------------------------------

@router.get("/{slug}/dossier")
def download_dossier(
    slug: str,
    audience: Literal["instructor", "student"] = "instructor",
    deployment_size: Literal["standard", "large"] = "standard",
    driver: Driver = Depends(get_driver),
):
    """GET /api/v1/attacks/{slug}/dossier — CYB-23 §12.2 / FR-LLM-06

    Resolves the URL slug to a KG ID, then:
      - If an LLM narrative has been pre-generated (via POST .../narrative),
        returns a prose .docx built from the stored text.
      - Otherwise, runs Q1–Q4 directly and returns a structured KG-data
        .docx that requires no LLM or API keys.
    """
    # Resolve slug → KG attack_id (e.g. 'colonial_pipeline_2021' → 'ATK-COL-001')
    try:
        attack_id = resolve_attack_id(slug)
    except ValueError:
        raise HTTPException(
            status_code=404,
            detail=(
                f"Attack '{slug}' not found. "
                "Valid slugs: colonial_pipeline_2021, triton_2017, "
                "german_steel_mill_2014, stuxnet_2010"
            ),
        )

    # ── Path 1: LLM narrative already stored ────────────────────────────────
    narrative = get_narrative(attack_id=attack_id, driver=driver)
    if narrative is not None:
        docx_bytes = _build_docx_from_narrative(
            slug=slug,
            narrative_text=narrative.narrative_text,
            audience=audience,
            deployment_size=deployment_size,
            model_used=narrative.narrative_model_used,
            generated_at=narrative.narrative_generated_at.isoformat(),
            kg_confidence=narrative.narrative_kg_confidence,
        )
        filename = f"{slug}_dossier.docx"
        return StreamingResponse(
            content=io.BytesIO(docx_bytes),
            media_type=DOCX_MIME,
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    # ── Path 2: KG-structured fallback (no LLM needed) ──────────────────────
    try:
        kg = _fetch_kg_data(driver=driver, attack_id=attack_id)
    except Exception as exc:
        raise HTTPException(
            status_code=503,
            detail=(
                f"Could not read KG data for '{slug}': {exc}. "
                "Ensure Neo4j is running and attack data has been loaded."
            ),
        ) from exc

    try:
        docx_bytes = _build_docx_from_kg(
            slug=slug,
            kg=kg,
            audience=audience,
            deployment_size=deployment_size,
        )
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to build dossier document: {exc}",
        ) from exc

    filename = f"{slug}_dossier.docx"
    return StreamingResponse(
        content=io.BytesIO(docx_bytes),
        media_type=DOCX_MIME,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# KG data fetcher
# ---------------------------------------------------------------------------

def _fetch_kg_data(driver: Driver, attack_id: str) -> dict:
    """Run Q1–Q4 and return a combined dict of KG data."""
    with driver.session() as s:
        # Q1: surface
        surface_rec = s.run(q1_surface.Q1_ATTACK_SURFACE, attack_id=attack_id).single()
        surface = dict(surface_rec) if surface_rec else {}

        # Q2: chain (one row per technique, ordered by step)
        chain_rows = [dict(r) for r in s.run(q2_chain.Q2_ATTACK_CHAIN, attack_id=attack_id)]

        # Q3: 4-layer consequence
        cons_rec = s.run(q3_consequence.Q3_CONSEQUENCE, attack_id=attack_id).single()
        cons = dict(cons_rec) if cons_rec else {}

        # Q4: AI/human roles
        roles_rec = s.run(q4_roles.Q4_ROLES, attack_id=attack_id).single()
        roles = dict(roles_rec) if roles_rec else {}

    return {
        "surface":      surface,
        "chain":        chain_rows,
        "layer_1_cyber":       list(cons.get("layer_1_cyber", [])),
        "layer_2_bridge":      list(cons.get("layer_2_bridge", [])),
        "layer_3_physical":    list(cons.get("layer_3_physical", [])),
        "layer_4_consequence": list(cons.get("layer_4_consequence", [])),
        "human_roles":         list(roles.get("human_roles", [])),
        "ai_components":       list(roles.get("ai_components", [])),
        "ai_attack_surfaces":  list(roles.get("ai_attack_surfaces", [])),
        "decision_points":     list(roles.get("decision_points", [])),
    }


# ---------------------------------------------------------------------------
# DOCX builder — Path 1: LLM narrative prose
# ---------------------------------------------------------------------------

def _build_docx_from_narrative(
    slug: str,
    narrative_text: str,
    audience: str,
    deployment_size: str,
    model_used: str,
    generated_at: str,
    kg_confidence: float,
) -> bytes:
    """Wrap the LLM narrative prose in a styled Word document."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    _set_margins(doc)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Attack Dossier — {slug.replace('_', ' ').title()}")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph()

    # Metadata table
    meta_rows = [
        ("Audience",        audience.capitalize()),
        ("Deployment Size", deployment_size.capitalize()),
        ("Generated At",    generated_at[:19].replace("T", " ") + " UTC"),
        ("Model Used",      model_used or "—"),
        ("KG Confidence",   f"{kg_confidence:.0%}"),
    ]
    _add_meta_table(doc, meta_rows)
    doc.add_paragraph()

    # Narrative body
    doc.add_heading("Narrative", level=1)
    for line in narrative_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=2)
        else:
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_after = Pt(6)

    _add_footer(doc)
    return _serialize(doc)


# ---------------------------------------------------------------------------
# DOCX builder — Path 2: KG-structured (no LLM)
# ---------------------------------------------------------------------------

def _build_docx_from_kg(
    slug: str,
    kg: dict,
    audience: str,
    deployment_size: str,
) -> bytes:
    """Build a structured teaching dossier directly from Q1–Q4 KG data."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    _set_margins(doc)

    surface = kg.get("surface", {})
    now_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    # ── Title ────────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Attack Dossier — {slug.replace('_', ' ').title()}")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("CyberKG-CPS  |  NSF Task 2.2  |  Instructor Edition")
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # ── Metadata table ───────────────────────────────────────────────────────
    meta_rows = [
        ("Attack",          slug.replace("_", " ").title()),
        ("Sector",          str(surface.get("industry_sector", "—"))),
        ("Year",            str(surface.get("year", "—"))),
        ("Attributed To",   str(surface.get("attributed_to", "—"))),
        ("Audience",        audience.capitalize()),
        ("Deployment Size", deployment_size.capitalize()),
        ("Generated At",    now_str + "  (KG-structured — no LLM)"),
    ]
    _add_meta_table(doc, meta_rows)
    doc.add_paragraph()

    # ── Section 1: Attack Overview ───────────────────────────────────────────
    doc.add_heading("1. Attack Overview", level=1)
    desc = surface.get("description") or surface.get("name") or "(no description in KG)"
    p = doc.add_paragraph(str(desc))
    p.paragraph_format.space_after = Pt(6)

    techniques = surface.get("techniques", [])
    if techniques:
        doc.add_heading("Techniques Used", level=2)
        for t in techniques:
            tech_id  = t.get("technique_id") or t.get("mitre_id", "")
            name     = t.get("name", "")
            tactic   = t.get("tactic", "")
            plane    = t.get("plane", "")
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(f"{tech_id}  {name}")
            r.bold = True
            if tactic or plane:
                p.add_run(f"  —  {tactic}  |  Plane: {plane}")

    # ── Section 2: IT-OT Attack Chain ───────────────────────────────────────
    chain = kg.get("chain", [])
    if chain:
        doc.add_heading("2. IT-OT Attack Chain", level=1)
        for step in chain:
            step_num = step.get("step", "—")
            tech_id  = step.get("technique_id") or step.get("mitre_id", "")
            name     = step.get("name", "")
            tactic   = step.get("tactic", "")
            plane    = step.get("plane", "")
            purdue   = step.get("purdue_level", "")
            phase    = step.get("phase", "")
            evidence = step.get("evidence_class", "")

            p = doc.add_paragraph(style="List Number")
            r = p.add_run(f"Step {step_num}  —  {tech_id}  {name}")
            r.bold = True
            details = "  |  ".join(filter(None, [
                f"Tactic: {tactic}" if tactic else "",
                f"Plane: {plane}" if plane else "",
                f"Purdue L{purdue}" if purdue else "",
                f"Phase: {phase}" if phase else "",
                f"Evidence: {evidence}" if evidence else "",
            ]))
            if details:
                p.add_run(f"\n    {details}")

    # ── Section 3: Consequence Chain ────────────────────────────────────────
    doc.add_heading("3. Consequence Chain", level=1)

    layer1 = kg.get("layer_1_cyber", [])
    layer2 = kg.get("layer_2_bridge", [])
    layer3 = kg.get("layer_3_physical", [])
    layer4 = kg.get("layer_4_consequence", [])

    if layer1:
        doc.add_heading("Layer 1 — Cyber (ATT&CK Techniques)", level=2)
        for item in layer1:
            name = item.get("name") or item.get("technique_id", "—")
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(name))

    if layer2:
        doc.add_heading("Layer 2 — Bridge Mechanism", level=2)
        for brg in layer2:
            name   = brg.get("name", "—")
            btype  = brg.get("bridge_type", "")
            pf     = brg.get("purdue_from", "")
            pt_    = brg.get("purdue_to", "")
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(str(name))
            r.bold = True
            if btype:
                p.add_run(f"  ({btype})")
            if pf or pt_:
                p.add_run(f"  |  Purdue {pf} → {pt_}")

    if layer3:
        doc.add_heading("Layer 3 — Physical Process", level=2)
        for item in layer3:
            name = item.get("name") or item.get("system_name", "—")
            purdue = item.get("purdue_level", "")
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(name))
            if purdue:
                p.add_run(f"  (Purdue L{purdue})")

    if layer4:
        doc.add_heading("Layer 4 — Physical Consequences", level=2)
        for c in layer4:
            name     = c.get("name", "—")
            cat      = c.get("table1_category") or c.get("category", "")
            severity = c.get("severity", "")
            realized = c.get("was_realized")
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(str(name))
            r.bold = True
            tags = "  |  ".join(filter(None, [
                f"Category: {cat}" if cat else "",
                f"Severity: {severity}" if severity else "",
                ("Realized" if realized else "Unrealized") if realized is not None else "",
            ]))
            if tags:
                p.add_run(f"  —  {tags}")

    # ── Section 4: AI / Human Roles ─────────────────────────────────────────
    human_roles  = kg.get("human_roles", [])
    ai_comps     = kg.get("ai_components", [])
    ai_surfaces  = kg.get("ai_attack_surfaces", [])
    dec_points   = kg.get("decision_points", [])

    doc.add_heading("4. AI / Human Roles", level=1)

    if human_roles:
        doc.add_heading("Human Actors", level=2)
        for h in human_roles:
            role = h.get("role_type") or h.get("role", "—")
            desc = h.get("role_description") or h.get("description", "")
            action = h.get("action", "")
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(str(role))
            r.bold = True
            if desc:
                p.add_run(f"  —  {desc}")
            if action:
                p.add_run(f"  |  Action: {action}")

    if ai_comps:
        doc.add_heading("AI Components", level=2)
        for ai in ai_comps:
            name = ai.get("name", "—")
            ai_type = ai.get("ai_type", "")
            hyp = ai.get("is_hypothetical")
            role = ai.get("role") or ai.get("scenario_description", "")
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(str(name))
            r.bold = True
            tags = "  |  ".join(filter(None, [
                str(ai_type) if ai_type else "",
                "Hypothetical" if hyp else "",
                str(role) if role else "",
            ]))
            if tags:
                p.add_run(f"  —  {tags}")

    if ai_surfaces:
        doc.add_heading("AI Attack Surfaces", level=2)
        for ais in ai_surfaces:
            name  = ais.get("name", "—")
            stype = ais.get("surface_type", "")
            atlas = ais.get("mitre_atlas_id", "")
            adv   = ais.get("is_adversarial")
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(str(name))
            r.bold = True
            tags = "  |  ".join(filter(None, [
                str(stype) if stype else "",
                str(atlas) if atlas else "",
                ("Adversarial" if adv else "Resilience") if adv is not None else "",
            ]))
            if tags:
                p.add_run(f"  —  {tags}")

    if dec_points:
        doc.add_heading("Decision Points", level=2)
        for dp in dec_points:
            desc = dp.get("description", "—")
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(desc))

    if not any([human_roles, ai_comps, ai_surfaces, dec_points]):
        doc.add_paragraph("No role data found in KG for this attack case.")

    _add_footer(doc)
    return _serialize(doc)


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def _set_margins(doc) -> None:
    from docx.shared import Inches
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)


def _add_meta_table(doc, rows: list[tuple[str, str]]) -> None:
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        row = tbl.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        for para in row.cells[0].paragraphs:
            for r in para.runs:
                r.bold = True


def _add_footer(doc) -> None:
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("Generated by CyberKG-CPS  |  NSF Task 2.2  |  Sprint 4 / T062")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x94, 0x96, 0x9A)
    run.italic = True


def _serialize(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
