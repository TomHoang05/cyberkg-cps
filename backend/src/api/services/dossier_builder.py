"""Dossier Builder — Sprint 4 / T062.

Converts a NarrativeResult into a styled .docx binary using python-docx.
Parses lightweight markdown from the LLM output (# h1, ## h2, ### h3,
- / * bullets, 1. numbered lists, **bold**, *italic*, blank-line paragraphs).

Public API
----------
build_dossier_docx(result: NarrativeResult) -> bytes
    Return a .docx file as an in-memory bytes object ready for StreamingResponse.
"""
from __future__ import annotations

import io
import re
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.api.services.narrative_service import NarrativeResult


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def build_dossier_docx(result: NarrativeResult) -> bytes:
    """Build a .docx from a NarrativeResult and return raw bytes."""
    doc = Document()

    # ── Page margins (1 in top/bottom, 1.25 in sides) ───────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # ── Title ────────────────────────────────────────────────────────────────
    attack_label = result.attack_id.replace("_", " ").title()
    title_para = doc.add_heading(f"{attack_label} — Attack Dossier", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Subtitle ─────────────────────────────────────────────────────────────
    subtitle = doc.add_paragraph("CyberKG-CPS  |  NSF Task 2.2  |  Instructor Edition")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ── Metadata line ────────────────────────────────────────────────────────
    generated_str = (
        result.narrative_generated_at.strftime("%Y-%m-%d %H:%M UTC")
        if isinstance(result.narrative_generated_at, datetime)
        else str(result.narrative_generated_at)
    )
    confidence_pct = f"{result.narrative_kg_confidence:.0%}" if result.narrative_kg_confidence else "N/A"
    meta = doc.add_paragraph(
        f"Generated: {generated_str}  |  "
        f"Model: {result.narrative_model_used or 'N/A'}  |  "
        f"KG Confidence: {confidence_pct}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Spacer before body
    doc.add_paragraph()

    # ── Body narrative ───────────────────────────────────────────────────────
    _render_markdown(doc, result.narrative_text or "")

    # ── Serialize to bytes ───────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Lightweight markdown renderer
# ---------------------------------------------------------------------------

def _render_markdown(doc: Document, text: str) -> None:
    """Parse lightweight markdown and add styled paragraphs to doc.

    Supported:
      # H1   ## H2   ### H3
      - / *  bullet items    1. numbered items
      **bold**  *italic*
      blank lines (ignored as spacers between paragraphs)
      everything else → normal paragraph (multi-line blocks joined)
    """
    lines = text.splitlines()
    i = 0
    prev_blank = True  # treat start-of-doc as after a blank

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Blank line — just advance
        if not stripped:
            prev_blank = True
            i += 1
            continue

        # H1
        if stripped.startswith("# ") and not stripped.startswith("## "):
            doc.add_heading(stripped[2:].strip(), level=1)
            prev_blank = False

        # H2
        elif stripped.startswith("## ") and not stripped.startswith("### "):
            doc.add_heading(stripped[3:].strip(), level=2)
            prev_blank = False

        # H3
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            prev_blank = False

        # Bullet list item (- or *)
        elif re.match(r'^[\-\*] ', stripped):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, stripped[2:].strip())
            prev_blank = False

        # Numbered list item
        elif re.match(r'^\d+[\.\)] ', stripped):
            p = doc.add_paragraph(style="List Number")
            body = re.sub(r'^\d+[\.\)] ', '', stripped).strip()
            _add_inline(p, body)
            prev_blank = False

        # Normal paragraph — collect consecutive non-special lines
        else:
            block_lines = [stripped]
            while i + 1 < len(lines):
                nxt = lines[i + 1].strip()
                if (
                    not nxt
                    or nxt.startswith("#")
                    or re.match(r'^[\-\*] ', nxt)
                    or re.match(r'^\d+[\.\)] ', nxt)
                ):
                    break
                i += 1
                block_lines.append(lines[i].strip())

            p = doc.add_paragraph()
            _add_inline(p, " ".join(block_lines))
            prev_blank = False  # noqa: F841

        i += 1


def _add_inline(para, text: str) -> None:
    """Add runs to para with **bold** and *italic* inline markers resolved."""
    # Split on **bold** and *italic* spans; keep delimiters via capturing group
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part:
            para.add_run(part)
